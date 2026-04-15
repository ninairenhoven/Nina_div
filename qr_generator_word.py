import io
import pandas as pd
import segno
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, Mm
from urllib.parse import urlparse


# KONSTANT
QR_WIDTH_MM = 25  # endre størrelse her


def is_valid_url_syntax(url: str) -> bool:
    try:
        parsed = urlparse(url)
        return all([parsed.scheme in ("http", "https"), parsed.netloc])
    except Exception:
        return False


def qr_png_bytes(url, scale=6, border=4):
    buf = io.BytesIO()
    segno.make(url, error='m').save(buf, kind='png', scale=scale, border=border)
    buf.seek(0)
    return buf

#====================================================================================================================
# Hjelpefunksjoner Word
#



def safe_name(s, max_len=100):
    s = str(s)
    return "".join(ch if ch.isalnum() or ch in " _-" else "_" for ch in s)[:max_len] or "uten_navn"


def fail_fast_if_not_writable(file_path: Path):
    file_path = Path(file_path)
    file_path.parent.mkdir(parents=True, exist_ok=True)

    if file_path.exists():
        # Test at filen kan åpnes for skriving uten å slette innholdet
        f = open(file_path, 'r+b')  # feiler hvis låst/ikke skrivetilgang
        f.close()
    else:
        # Filen finnes ikke: test at den kan opprettes
        f = open(file_path, 'xb')   # feiler hvis ikke kan opprettes
        f.close()
        file_path.unlink()          # rydd opp



def set_document_margins(doc, top=25, right=30, bottom=25, left=30):
    for section in doc.sections:
        section.top_margin = Mm(top)
        section.right_margin = Mm(right)
        section.bottom_margin = Mm(bottom)
        section.left_margin = Mm(left)



#====================================================================================================================
# Lage URL og skrive til word
#

# Felles skrivelogikk for én gruppe
def write_qr_group(doc, title, urls, subtitle='', spacing=12):
    # Overskrift for gruppen
    h = doc.add_heading(level=1)
    r = h.add_run(str(title))
    r.font.size = Pt(18)
    if (subtitle!=''):
        h2 = doc.add_heading(level=2)
        r = h2.add_run(str(subtitle))
        r.font.size = Pt(12)
    r.add_break()
    #
    # Innhold: URL-linje + QR på ny linje
    for idx, url in enumerate(urls):
        p = doc.add_paragraph()
        p.add_run().add_picture(qr_png_bytes(url), width=Mm(QR_WIDTH_MM))
        p.add_run('\t'+str(url))
        
        if spacing>0 and idx < len(urls) - 1:
            # Space mellom koder
            #print(spacing)
            p = doc.add_paragraph()
            #p.add_run(str(spacing))
            p.paragraph_format.space_after = Pt(spacing)
            #r = p.add_run()
            #r.add_break()  # linjeskift



# ÉN samlet fil med alle grupper
def export_qr_one_file(df, out_folder, group="Tiltak", main_title="", page_breaks=True, spacing=12):
    """
    df må ha kolonnene: group og URL (trimmet/filtrert).
    Lager én .docx med hovedtittel og en seksjon per group-verdi.
    Returnerer filstien.
    """
    out_folder = Path(out_folder)
    out_folder.mkdir(parents=True, exist_ok=True)
    out_path = out_folder / f"{safe_name(main_title)}.docx"
    #
    fail_fast_if_not_writable(out_path)  # feiler tidlig hvis ikke mulig
    #
    # Stabil rekkefølge
    df_sorted = df.sort_values([group, "URL"])
    doc = Document()
    set_document_margins(doc)
    #
    if main_title:
        h0 = doc.add_heading(level=0)
        rh0 = h0.add_run(str(main_title))
        rh0.font.size = Pt(22)
    #
    for grp_value, grp_df in df_sorted.groupby(group, sort=False):
        print(f"{grp_value:<30} ({grp_df.index.size})")
        urls = grp_df["URL"].tolist()
        subtitle = '' #f'Antall: {grp_df.index.size}'
        write_qr_group(doc, title=grp_value, urls=urls, subtitle=subtitle, spacing=spacing)    
        if page_breaks:
            doc.add_page_break()
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(spacing)
    #
    doc.save(out_path)
    print('Lagret til '+str(out_path))
    return out_path


# ÉN fil per gruppe
def export_qr_per_group(df, out_folder, common_title="", group="Tiltak", spacing=12):
    """
    df må ha kolonnene: group og URL (trimmet/filtrert).
    Lager én .docx per group-verdi.
    Returnerer liste over filstier.
    """
    out_folder = Path(out_folder)
    out_folder.mkdir(parents=True, exist_ok=True)

    df_sorted = df.sort_values([group, "URL"])
    paths = []

    for grp_value, grp_df in df_sorted.groupby(group, sort=False):
        print(f"{grp_value:<30} ({grp_df.index.size})", end=' ')

        out_path = out_folder / f"{safe_name(grp_value)}.docx"
        fail_fast_if_not_writable(out_path)  # feiler tidlig hvis ikke mulig

        urls = grp_df["URL"].tolist()
        doc = Document()
        set_document_margins(doc)
        # Skriv innhold for denne gruppen
        write_qr_group(doc, title=grp_value, subtitle=common_title, urls=urls, spacing=spacing)

        doc.save(out_path)
        paths.append(out_path)
        print(out_path)
    print(f"Lagret {len(paths)} filer")
    return paths



#====================================================================================================================
# Lese excel-fil med lenker
#

def read_links_excel(file):
    df = pd.read_excel(file, dtype=str)
    df = df.loc[:, ~df.columns.str.startswith("Unnamed")]
    #
    df["Tiltak"] = df["Tiltak"].str.strip()
    df = df[df["Lenke"].notna() & (df["Lenke"] != "")]
    df = df.rename(columns={'Lenke':'URL'})
    #
    print('\nLeser fil: '+str(file))
    print(df)
    #
    df['valid'] = df['URL'].apply(is_valid_url_syntax)
    #
    firstpart = df['URL'].str.rsplit('/', n=1).str[0]
    print('\nURL frem til siste slash:')
    print(firstpart.value_counts())
    #
    if  (~df['valid'].all()):
        print('ADVARSEL: invalid url')
        print(df.loc[~df['valid']])
    else:
        print('Alle URL har gyldig syntaks')
    
    if df['URL'].duplicated().any():
        print('ADVARSEL: dupliserte lenker')
        print(df.loc[df['URL'].duplicated(keep=False)])
    else:
        print('Alle URL unike')
    return df

#====================================================================================================================


path_user = Path.home()
path = path_user.joinpath('OneDrive - Opinion AS/Documents/QR generator')

EXCEL1  = path.joinpath('BTU høst 2025_Lenker til tiltakene_TEST.xlsx')
EXCEL2  = path.joinpath('BTU høst 2025_Lenker til tiltakene_til Sissel.xlsx')
OUT  = path.joinpath('output')


df1 = read_links_excel(EXCEL1)
df2 = read_links_excel(EXCEL2)


"""
_ = export_qr_one_file(df1, OUT, group = 'Tiltak', main_title = 'QR-koder for brukerundersøkelsen, testlenker', page_breaks=True, spacing=120)

_ = export_qr_per_group(df2, OUT, group = 'Tiltak', common_title='QR-koder for brukerundersøkelsen', spacing=120)

"""

#====================================================================================================================

# Vise en QR på skjermen
"""
url = 'https://survey.opinion.no/oJuXlV/9zsk8a'

import tkinter as tk
from PIL import Image, ImageTk

img = Image.open(qr_png_bytes(url))
root = tk.Tk()
root.title("QR")
photo = ImageTk.PhotoImage(img)
tk.Label(root, image=photo).pack()
root.mainloop()

"""