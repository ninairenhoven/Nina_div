import pandas as pd
from docx import Document
import os
from pathlib import Path
from datetime import datetime, timedelta
import logging
import shutil
import gc
from xhtml2pdf import pisa
from concurrent.futures import ProcessPoolExecutor, as_completed
import multiprocessing
#import time
import re
from tqdm import tqdm  # Progress bar
import zipfile
from tkinter.filedialog import askdirectory
from tkinter import Tk

# Mammoth import moved to the top to avoid repeated imports
try:
    from mammoth import convert_to_html
    MAMMOTH_AVAILABLE = True
except ImportError:
    MAMMOTH_AVAILABLE = False
    print("❗ Warning: mammoth not installed, using basic conversion")

"""
================================================================================
DIGIPOST BREV-GENERATOR - BRUKERVEILEDNING
================================================================================

FORMÅL:
    Massegenerering av personaliserte PDF-brev fra Word/HTML-mal + Excel/CSV-fil.
    Støtter parallell prosessering og batch-organisering.

FORBEREDELSER:
    1. Lag brevmal med plassholdere i formatet [placeholder], f.eks.:
       - [name]
       - [uniksurveylink]
       - [mic]
       
       ALTERNATIV A: Word-mal (.docx)
       - Scriptet konverterer automatisk til HTML
       - Støtter header og footer fra Word-dokumentet
       - Støtter ikke bilder
       
       ALTERNATIV B: HTML-mal (.html)
       - Full kontroll over formatering og styling
       - Raskere (ingen konvertering nødvendig)
       - Krever HTML/CSS-kunnskap
       - Mulig å embedde logo
    
    2. Lag Excel/CSV-fil med kolonner som matcher plassholderne
       Pluss:
       - Kolonne med filnavn for hver PDF (f.eks. 'filnavn')
       - Kolonne med batch-nummer (f.eks. 'batch')
    
    3. Konfigurer stiene i INPUT-seksjonen:
       - WORD_TEMPLATE_FILE: Sti til Word-mal (hvis du bruker Word)
       - HTML_TEMPLATE_FILE: Sti til HTML-mal (hvis du bruker HTML)
       - LOOKUP_FILE: Excel/CSV-fil med mottakerdata
       - OUTPUT_FOLDER: Genereres automatisk med timestamp
       - PLACEHOLDERS: Liste med plassholdere (må matche mal + lookup-fil)
       - FILENAME_COL: Kolonnenavn for PDF-filnavn
       - BATCH_COL: Kolonnenavn for batch-inndeling

SPESIALTEGN I WORD-MAL:
    - <linjeskift>: Legger til ekstra linjeskift i PDF
    - E-postadresser: Konverteres automatisk til klikkbare mailto-lenker
    - URL-er: Formateres automatisk som klikkbare lenker

CSS-TILPASNING:
    (Gjelder kun hvis du genererer fra Word - hvis du bruker ferdig HTML-fil,
     må CSS-endringer gjøres direkte i HTML-filen)
    Juster typografi og layout i seksjonen "HTML CONFIGURATION" (linje ~75):
    - Fonter, fontstørrelser, linjeavstand
    - Marginer (topp/bunn/venstre/høyre)
    - Header og footer-styling

KJØRING:
    1. Kjør: python digipost_brev_generator_FINAL.py
    
    2. Svar på spørsmålene:
       a) Antall rader:
          [Y] = Les hele filen eller [tall] = Les kun X antall rader
       
       b) Batch-valg:
          [A] = Alle batcher eller [R] = Spesifiser range (f.eks. "1-5")
       
       c) PDF-generering:
          [G] = Generer nye PDF-filer eller [E] = Bruk eksisterende mappe (bare zipping)
       
       d) Template-kilde:
          [G] = Generer HTML fra Word-fil (anbefalt hvis du har Word-mal)
          [H] = Les ferdig HTML-fil (anbefalt hvis du har HTML-mal)
       
       e) Zipping:
          [Y] = Lag zip-filer per batch eller [N] = Hopp over zipping

OUTPUT:
    - OUTPUT_FOLDER/: PDF-filer (én per rad i lookup-filen)
    - OUTPUT_FOLDER/generation_log_TIMESTAMP.log: Loggfil
    - OUTPUT_FOLDER/html_template.html: Generert HTML-mal (for inspeksjon)
    - OUTPUT_FOLDER/BatchXX_YYYrows.zip: Zip-filer per batch (valgfritt)

TESTMODUS:
    Rask testing med 2 rader:
    python digipost_brev_generator_FINAL.py test        (Word-mal)
    python digipost_brev_generator_FINAL.py test html   (HTML-mal)

ARBEIDSFLYT - ANBEFALING:
    1. Start med Word-mal (enklere å redigere)
    2. Generer noen test-PDF-er
    3. Sjekk html_template.html i output-mappen
    4. Hvis du trenger finere CSS-kontroll: Rediger html_template.html
    5. Sett HTML_TEMPLATE_FILE til den redigerte filen
    6. Kjør på nytt med [H] for å bruke HTML-malen direkte

YTELSE:
    - Bruker multiprocessing (alle CPU-kjerner minus 1)
    - Prosesserer i chunks på 2000 filer
    - Viser estimert tid etter første chunk
    - For best ytelse: Bruk SSD for OUTPUT_FOLDER

FEILSØKING:
    - Sjekk loggfilen for detaljer om feil
    - Verifiser at alle plassholdere finnes både i mal og lookup-fil
    - Sjekk at FILENAME_COL og BATCH_COL matcher kolonnenavn i lookup-fil
    - Test med få rader først (2-10 rader) før full kjøring

AVHENGIGHETER:
    pip install pandas openpyxl python-docx xhtml2pdf mammoth tqdm

================================================================================
"""

# =============================================================================
# INPUT
# =============================================================================

# Define paths for user and input files
user_path = Path.home()
path = user_path.joinpath(r'Documents\RVU_LOKAL\BRAKAR Sample')

WORD_TEMPLATE_FILE = path.joinpath("Invitasjonsbrev_Brakar_Opinion Reisedag 18.11.docx")
HTML_TEMPLATE_FILE = path.joinpath("BRAKAR_TEMPLATE_Opinion_enzo_embedded.html")
#HTML_TEMPLATE_FILE = path.joinpath("BRAKAR_TEMPLATE_Brakar_logo_embedded.html")

#LOOKUP_FILE = path.joinpath("UTTREKK/Buskerud FK, Brakar, Opinion FREG uttrekk-krr_BRAKAR_Batch_1_13_PROCESSED_Nov14.csv")
LOOKUP_FILE = path.joinpath("UTTREKK/Buskerud FK, Brakar, Opinion FREG uttrekk-krr_BRAKAR_Batch_1_13_PROCESSED_Nov14_EDIT_Nov25.csv")

timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
OUTPUT_FOLDER = path.joinpath(f"OUTPUT_{timestamp}")

# IMPORTANT: For best performance, use SSD for OUTPUT_FOLDER!

# Define placeholders to be replaced in the template
PLACEHOLDERS = [
#    "mic",
    "name",
    "uniksurveylink"
]

# Column with file name for individual pdf file
FILENAME_COL = 'filnavn'

# Column defining which files will be zipped together
BATCH_COL = 'batch'

# =============================================================================
# HTML CONFIGURATION
# =============================================================================

# Default styles for the generated HTML
DEFAULT_FONT_FAMILY = "Helvetica, Arial, sans-serif"
DEFAULT_FONT_SIZE = "12pt"
DEFAULT_LINE_HEIGHT = "1.3"
LIST_LINE_HEIGHT = "1.15"
HEADING_1_SIZE = "14pt"
HEADING_2_SIZE = "12pt"
LINK_FONT_SIZE = "11.5pt"

MARGIN_TOP = "25mm"
MARGIN_BOTTOM = "15mm"
MARGIN_LEFT = "25mm"
MARGIN_RIGHT = "25mm"

# Header Configuration
HEADER_FONT_SIZE = "11.5pt"
HEADER_COLOR = "#808080"  # Grå farge (hex-kode)
HEADER_ALIGNMENT = "center"  # Midtstilt

# Footer Configuration
FOOTER_FONT_SIZE = "11.5pt"
FOOTER_COLOR = "#808080"  # Grå farge (hex-kode)
FOOTER_ALIGNMENT = "center"  # Midtstilt


# HTML template for the generated PDF
from string import Template

HTML_TEMPLATE = Template("""
<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <style>
        @page {
            size: A4;
            margin-top: $MARGIN_TOP; 
            margin-bottom: $MARGIN_BOTTOM; 
            margin-left: $MARGIN_LEFT; 
            margin-right: $MARGIN_RIGHT;
        }
        body {
            font-family: $DEFAULT_FONT_FAMILY;
            font-size: $DEFAULT_FONT_SIZE;
            line-height: $DEFAULT_LINE_HEIGHT;
            color: #000;
        }
        h1, h2, h3 {
            color: #000;
            margin-top: 1.5em;
            margin-bottom: 0.6em;
        }
        h1 { font-size: $HEADING_1_SIZE; font-weight: bold; }
        h2 { font-size: $HEADING_2_SIZE; font-weight: bold; }
        p {
            margin: 0.5em 0;
        }
        strong {
            font-weight: bold;
        }
        ul, ol {
            margin: 0.5em 0;
            padding-left: 2em;
            line-height: $LIST_LINE_HEIGHT;
        }
        li {
            margin: 0.3em 0;
            line-height: $LIST_LINE_HEIGHT;
        }
        a.survey-link {
            color: #0066cc;
            text-decoration: underline;
            font-size: $LINK_FONT_SIZE;
        }
        a.email-link {
            color: #0066cc;
            text-decoration: underline;
        }
    </style>
</head>
<body>
<header style="text-align: $HEADER_ALIGNMENT; font-size: $HEADER_FONT_SIZE; color: $HEADER_COLOR; font-family: $DEFAULT_FONT_FAMILY; margin: 0.5em 0;">
    $HEADER_CONTENT
</header>
$HTML_BODY
<footer>
    <div style="
        text-align: center; 
        font-size: $FOOTER_FONT_SIZE; 
        color: $FOOTER_COLOR; 
        font-family: $DEFAULT_FONT_FAMILY; 
        margin: 0.5em 0;">
        $FOOTER_CONTENT
    </div>
</footer>
</body>
</html>
""")


# =============================================================================
# LOAD HTML TEMPLATE
# =============================================================================


def extract_header_footer(template_path):
    """Extract header and footer content from a Word document."""
    try:
        logging.info(f"Extracting header and footer from: {template_path}")
        doc = Document(template_path)
        
        # Hent header-innhold
        headers = []
        for section in doc.sections:
            header = section.header
            for para in header.paragraphs:
                if para.text.strip():
                    headers.append(f"<header>{para.text}</header>")
        
        header_content = "\n".join(headers)

        # Hent footer-innhold
        footers = []
        for section in doc.sections:
            footer = section.footer
            for para in footer.paragraphs:
                if para.text.strip():
                    footers.append(f"<footer>{para.text}</footer>")
        
        footer_content = "\n".join(footers)

        logging.info("Successfully extracted header and footer")
        return header_content, footer_content

    except Exception as e:
        logging.error(f"Failed to extract header and footer: {e}")
        raise Exception(f"Error extracting header and footer: {e}")



def generate_html_template(template_path):
    """Load and convert template to HTML"""
    try:
        logging.info(f"Loading template: {template_path}")

        if MAMMOTH_AVAILABLE:
            logging.info("Using Mammoth for HTML conversion")
            with open(template_path, "rb") as docx_file:
                result = convert_to_html(docx_file)
                body = result.value
        else:
            logging.info("Using basic conversion (Mammoth not available)")
            doc = Document(template_path)
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text
                if text.strip():
                    paragraphs.append(f"<p>{text}</p>")
            body = "\n".join(paragraphs)
            logging.info(f"Processed {len(paragraphs)} paragraphs")
        
        # Process spacing tags and emails
        body = convert_emails_to_links(body)
        body = process_spacing_tags(body)
        
        for placeholder in PLACEHOLDERS:
            if f'[{placeholder}]' not in body:
                logging.warning(f"❗ Placeholder [{placeholder}] not found in template!")

        logging.info("Template successfully converted to HTML")

        # Hent header og footer
        header_content, footer_content = extract_header_footer(template_path)

        # Kombiner header, body og footer
        # body = f"{header_content}\n{body}\n{footer_content}"

        return HTML_TEMPLATE.substitute(            
            MARGIN_TOP = MARGIN_TOP,
            MARGIN_BOTTOM = MARGIN_BOTTOM,
            MARGIN_LEFT = MARGIN_LEFT,
            MARGIN_RIGHT = MARGIN_RIGHT,
            DEFAULT_FONT_FAMILY=DEFAULT_FONT_FAMILY,
            DEFAULT_FONT_SIZE=DEFAULT_FONT_SIZE,
            DEFAULT_LINE_HEIGHT=DEFAULT_LINE_HEIGHT,
            LIST_LINE_HEIGHT=LIST_LINE_HEIGHT,
            HEADING_1_SIZE=HEADING_1_SIZE,
            HEADING_2_SIZE=HEADING_2_SIZE,
            LINK_FONT_SIZE=LINK_FONT_SIZE,
            HEADER_FONT_SIZE=HEADER_FONT_SIZE,
            HEADER_COLOR=HEADER_COLOR,
            HEADER_ALIGNMENT=HEADER_ALIGNMENT,
            FOOTER_FONT_SIZE=FOOTER_FONT_SIZE,
            FOOTER_COLOR=FOOTER_COLOR,
            FOOTER_ALIGNMENT=FOOTER_ALIGNMENT,
            HEADER_CONTENT=header_content,
            FOOTER_CONTENT=footer_content,
            HTML_BODY=body
        )

    except Exception as e:
        logging.error(f"Failed to load template: {e}")
        raise Exception(f"Failed to load template in worker: {e}")




def convert_emails_to_links(html_content):
    """Convert email addresses in HTML to clickable mailto: links"""
    email_pattern = r'\b([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})\b'
    
    def replace_email(match):
        email = match.group(1)
        return f'<a href="mailto:{email}" class="email-link">{email}</a>'
    
    return re.sub(email_pattern, replace_email, html_content)


def process_spacing_tags(html_content):
    """Process spacing tags from Word document"""
    tag_count = html_content.count('&lt;linjeskift&gt;')
    if tag_count > 0:
        html_content = html_content.replace('&lt;linjeskift&gt;', '<p>&nbsp;</p>')
    return html_content



# =============================================================================
# READ LOOKUP FILE
# =============================================================================

def load_and_validate_data(file_path, required_columns, limit_rows=None):
    """
    Load and validate data from a lookup file (Excel or CSV).

    Parameters:
    -----------
    file_path : str or Path
        Path to the lookup file (Excel or CSV).
    required_columns : list
        List of required column names.
    limit_rows : int, optional
        Limit number of rows to read (for testing).

    Returns:
    --------
    pd.DataFrame
        Loaded and validated data as a Pandas DataFrame.

    Raises:
    -------
    ValueError
        If the file is missing required columns or cannot be loaded.
    """
    file_path = Path(file_path).resolve()
    file_extension = file_path.suffix.lower()

    try:
        # Load the file based on its extension
        if file_extension in ['.xlsx', '.xls']:
            df = pd.read_excel(
                file_path,
                usecols=required_columns,
                nrows=limit_rows,
                engine='openpyxl',
                #dtype={'filename_tag': str}  # Preserve leading zeros
            )
        elif file_extension == '.csv':
            df = pd.read_csv(
                file_path,
                usecols=required_columns,
                nrows=limit_rows,
                #dtype={'filename_tag': str}  # Preserve leading zeros
            )
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
    except ValueError as e:
        raise ValueError(f"Failed to load file: {e}")

    # Validate required columns
    missing_columns = [col for col in required_columns if col not in df.columns]
    if missing_columns:
        raise ValueError(f"Missing required columns: {missing_columns}")
    
    # Check for missing values in required columns
    for col in required_columns:
        missing_values_count = df[col].isna().sum() + df[col].eq('').sum()
        if missing_values_count > 0:
            raise ValueError(f"Column '{col}' has {missing_values_count} missing values!")

    batch_counts = df[BATCH_COL].value_counts().sort_index()
    print(f'\n{len(df)} rader, {len(batch_counts)} batch(er):\n')
    print(df)

    logging.info(f"Loaded {len(df)} rows from {file_path}")
    logging.info(f"Columns: {list(df.columns)}")
    
    return df




# =============================================================================
# GLOBAL WORKER TEMPLATE CACHE - Optimization for multiprocessing
# =============================================================================
_worker_template_cache = None


def init_worker(html_template):
    """
    Initialize worker process with template (called once per worker).
    This saves memory usage by avoiding sending the template via pickle to each worker.
    """
    global _worker_template_cache
    if _worker_template_cache is None:
        _worker_template_cache = html_template


def process_single_file_worker(row_dict, output_folder, required_placeholders):
    """
    Worker function for parallel processing.
    Uses cached template from worker process.
    """
    global _worker_template_cache
    
    try:
        # Use cached template
        content = _worker_template_cache
        
        # Replace placeholders
        for placeholder in required_placeholders:
            value = str(row_dict[placeholder])#str(row_dict.get(placeholder, ''))
            
            # Hvis verdien starter med "http://" eller "https://", formater den som lenke
            if value.startswith("http://") or value.startswith("https://"):
                value = f'<a href="{value}" class="survey-link">{value}</a>'
            
            content = content.replace(f'[{placeholder}]', value)
        
        # Generate filename
        # filename_tag = str(row_dict['filename_tag'])
        # filename_tag_clean = filename_tag.replace('/', '_').replace('\\', '_').replace(':', '_')
        filename = row_dict[FILENAME_COL] #f"brakar_opinion_{filename_tag_clean}.pdf"
        pdf_path = os.path.join(output_folder, filename)
        
        # Create PDF using xhtml2pdf
        with open(pdf_path, 'wb') as pdf_file:
            pisa_status = pisa.CreatePDF(
                content.encode('utf-8'),
                dest=pdf_file,
                encoding='utf-8'
            )
        
        if pisa_status.err:
            return {'success': False, 'error': 'xhtml2pdf conversion failed'}
        
        return {'success': True, 'filename': filename}
    
    except Exception as e:
        filename = row_dict.get(FILENAME_COL, 'unknown')
        return {'success': False, 'error': str(e)}

# =============================================================================
# GENERATOR
# =============================================================================

class HTMLPDFGenerator:
    """
    Fast HTML-based PDF generator with parallel processing using xhtml2pdf
    """
    def __init__(self, lookup, output_folder, required_placeholders, html_template, chunk_size=2000, num_workers=None):
        """
        Initialize generator

        Parameters:
        -----------
        html_template : str, optional
            Pre-generated HTML template as a string. If provided, this will be used.
        template_path : str or Path, optional
            Path to a Word template. Used only if html_template is not provided.
        lookup : pd.DataFrame
            Preloaded and validated DataFrame with data to merge into template
        output_folder : str or Path
            Output folder for PDFs
        required_placeholders : list
            List of required placeholder names (without [])
        chunk_size : int
            Number of files to process per chunk
        num_workers : int
            Number of parallel workers (default: CPU count - 1)
        """
        self.html_template = html_template
        
        #self.template_path = str(Path(template_path).resolve())
        self.lookup = lookup  # DataFrame is now passed in directly
        self.output_folder = str(Path(output_folder).resolve())
        self.required_placeholders = required_placeholders
        self.chunk_size = chunk_size
        self.num_workers = num_workers or max(1, multiprocessing.cpu_count() - 1)

        # Create output folder
        Path(self.output_folder).mkdir(parents=True, exist_ok=True)

        # Cleanup old temp folders
        self.cleanup_temp_folders()

        # Get root logger (inherits your basicConfig setup)
        self.logger = logging.getLogger()

        # Stats
        self.stats = {
            'total': len(lookup),
            'success': 0,
            'failed': 0,
            'start_time': None,
            'end_time': None
        }

    def cleanup_temp_folders(self):
        """Remove old temp folders from previous runs"""
        for folder in Path(self.output_folder).glob("temp_*"):
            if folder.is_dir():
                try:
                    shutil.rmtree(folder)
                except:
                    pass

    def process_chunk(self, chunk_df, chunk_num, output_folder):
        """
        Process a chunk of files in parallel with worker template caching
        """
        self.logger.info(f"\n\nProcessing chunk {chunk_num} ({len(chunk_df)} files)...")
        t0 = datetime.now()

        chunk_success = 0
        chunk_failed = 0

        # Optimize: Workers initialized with template via initializer
        with ProcessPoolExecutor(
            max_workers=self.num_workers,
            initializer=init_worker,
            initargs=(self.html_template,)
        ) as executor:
            # Submit all tasks
            futures = {}
            for idx, row in chunk_df.iterrows():
                # Send only necessary data (not the entire row)
                row_dict = {
                    col: row[col]
                    for col in self.required_placeholders + [FILENAME_COL] #+ ['filename_tag']
                }

                future = executor.submit(
                    process_single_file_worker,
                    row_dict,
                    output_folder,
                    self.required_placeholders
                )
                futures[future] = row.get(FILENAME_COL, idx) #row.get('filename_tag', idx)

            # Collect results with progress bar
            with tqdm(total=len(futures), desc=f"Chunk {chunk_num}") as pbar:
                for future in as_completed(futures):
                    filename = futures[future]
                    try:
                        result = future.result()
                        if result['success']:
                            chunk_success += 1
                            self.stats['success'] += 1
                        else:
                            chunk_failed += 1
                            self.stats['failed'] += 1
                            self.logger.error(f"Failed {filename}: {result.get('error', 'Unknown')}")
                    except Exception as e:
                        chunk_failed += 1
                        self.stats['failed'] += 1
                        self.logger.error(f"Failed {filename}: {str(e)}")
                    finally:
                        pbar.update(1)  # Update progress bar
        
            chunk_duration = (datetime.now() - t0).total_seconds()

            self.logger.info(f"Chunk {chunk_num} complete: {chunk_success} success, {chunk_failed} failed, {chunk_duration/60:.1f} min")
            gc.collect()

            return chunk_success, chunk_failed

    def generate(self):
        """
        Main generation process
        """
        self.stats['start_time'] = datetime.now()
        self.logger.info(f"{'='*68}")
        self.logger.info(f"Starting PDF Generation (xhtml2pdf + worker caching)")
        self.logger.info(f"{'='*68}")
        self.logger.info(f"Output: {self.output_folder}")
        self.logger.info(f"Parallel workers: {self.num_workers}")
        self.logger.info(f"Chunk size: {self.chunk_size}")

        # Process in chunks
        num_chunks = (len(self.lookup) + self.chunk_size - 1) // self.chunk_size
        self.logger.info(f"Number of chunks: {num_chunks}")

        chunk_start_time = None
        files_remaining = len(self.lookup)

        for chunk_num in range(num_chunks):
            #if chunk_num == 0:
            chunk_start_time = datetime.now()
            start_idx = chunk_num * self.chunk_size
            end_idx = min((chunk_num + 1) * self.chunk_size, len(self.lookup))
            chunk_df = self.lookup.iloc[start_idx:end_idx]

            self.process_chunk(chunk_df, chunk_num + 1, self.output_folder)

            # Estimate remaining time
            #if chunk_num == 0 and num_chunks > 1:
            chunk_duration = (datetime.now() - chunk_start_time).total_seconds()
            files_processed = len(chunk_df)
            files_remaining = files_remaining - files_processed
            
            # Estimate based on first chunk
            est_remaining_seconds = (chunk_duration / files_processed) * files_remaining
            est_finish_time = datetime.now() + timedelta(seconds=est_remaining_seconds)
            
            #self.logger.info(f"\nTIME ESTIMATE:")
            #self.logger.info(f"    ")
            self.logger.info(f"⏱️ {files_remaining} files remaining. Estimated {est_remaining_seconds/60:.1f} min, done at ~{est_finish_time.strftime('%H:%M:%S')}")
            #self.logger.info(f"   Estimated completion: {est_finish_time.strftime('%H:%M:%S')}")
            
            #print(f"\n  ESTIMATE: {est_remaining_seconds/60:.1f} min remaining, done at ~{est_finish_time.strftime('%H:%M:%S')}\n")

        # Final stats
        self.stats['end_time'] = datetime.now()
        duration = (self.stats['end_time'] - self.stats['start_time']).total_seconds()
        self.logger.info(f"{'='*68}")
        self.logger.info(f"✅ GENERATION COMPLETE")
        self.logger.info(f"{'='*68}")
        self.logger.info(f"Total files: {self.stats['total']}")
        self.logger.info(f"Successful: {self.stats['success']}")
        if self.stats['failed']>0:
            self.logger.info(f"❌ Failed: {self.stats['failed']}")
        else:
            self.logger.info(f"Failed: {self.stats['failed']}")
        self.logger.info(f"Total time: {duration/60:.1f} minutes")
        self.logger.info(f"{'='*68}")

        return True


# =============================================================================
# LOGGER
# =============================================================================


def setup_logging(output_folder):
    """
    Setup logging to file and console
    Should be called before any processing starts
    
    Parameters:
    -----------
    output_folder : Path
        Folder where log file will be created
    
    Returns:
    --------
    logger : logging.Logger
        Configured logger instance
    """
    log_filename = output_folder.joinpath(
        f'generation_log_{datetime.now().strftime("%Y%m%d_%H%M%S")}.log'
    )
    
    # Clear any existing handlers to avoid duplicates
    for handler in logging.root.handlers[:]:
        logging.root.removeHandler(handler)
    
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(levelname)s - %(message)s',
        handlers=[
            logging.FileHandler(log_filename, encoding='utf-8'),
            logging.StreamHandler()
        ]
    )
    
    logger = logging.getLogger(__name__)
    logger.info(f"Log file created: {log_filename}")
    
    return logger




#lookup_df = load_and_validate_data(file_path=LOOKUP_FILE, required_columns=PLACEHOLDERS + [FILENAME_COL], limit_rows = None) #['filename_tag'], 



# =============================================================================
# HELPER FUNCTIONS FOR MAIN
# =============================================================================

def get_row_limit():
    """
    Ask user for row limit for lookup file.
    
    Returns:
        int or None: Number of rows to process, or None for all rows
    """
    user_input = input("\n❓ Read entire lookup file [Y] or specify number of rows [digit]: ").strip().lower()

    if user_input == "y":
        print("\n✅ Read entire lookup file...")
        return None
    elif user_input.isdigit():
        limit_rows = int(user_input)
        print(f"\n✅ Read {limit_rows} rows")
        return limit_rows
    else:
        print("\n❌ Invalid input")
        exit(1)


def generate_pdfs(template, lookup_df):
    """
    Generate PDF files from template and lookup data.
    
    Parameters:
        template (str): HTML template for PDF generation
        lookup_df (pd.DataFrame): Lookup data
        placeholders (list): List of placeholder names
    
    Returns:
        Path: Path to output folder with generated PDFs
    """
    # Create output folder
    Path(OUTPUT_FOLDER).mkdir(parents=True, exist_ok=True)
    
    # Initialize logging
    logger = setup_logging(OUTPUT_FOLDER)
    
    # Save HTML template
    html_output_path = OUTPUT_FOLDER.joinpath("html_template.html")
    with open(html_output_path, "w", encoding="utf-8") as html_file:
        html_file.write(template)
    print(f"HTML-templatet er lagret til: {html_output_path}")
    
    # Generate PDFs
    generator = HTMLPDFGenerator(
        html_template=template,
        lookup=lookup_df,
        output_folder=str(OUTPUT_FOLDER),
        required_placeholders=PLACEHOLDERS,
        chunk_size=2000,
        num_workers=None  # Auto-detect CPU cores
    )
    generator.generate()
    
    return OUTPUT_FOLDER


def select_batches(df):
    """
    Ask user which batches to include.
    #
    Returns:
        df with selected batches
    """
    batch_choice = input("\n❓ Include all batches [A (default)] or specify range [R]? ").strip().lower()
    #
    if batch_choice in ["a", ""]:
        print("\n✅ Include all batches...")
        use_df = df
        #return df
    elif batch_choice == "r":
        try:
            selection = input("Include batch range [from-to]:\t").split('-')
            selection = [int(x) for x in selection]
            
            if len(selection) == 2:
                #from_batch, to_batch = batch_selection
                batch_range = range(selection[0], selection[1]+1)

            elif len(selection) == 1:
                #from_batch, to_batch = batch_selection[0], batch_selection[0]
                batch_range = range(selection[0], selection[0]+1)
            else: 
                print("\n❌ Invalid batch input")
                exit(1)

            mask = df[BATCH_COL].isin(batch_range)
            print(f"\n✅ Include {len(batch_range)} batch(es), {mask.sum()} rows\n")
            use_df = df.loc[mask]
            #return df.loc[mask]
        except ValueError:
            print("\n❌ Invalid batch input")
            exit(1)

    if len(use_df)>10000:
        check_input = input('❗ WARNING: More Than 10 000 rows. [Y] to continue: ').lower()
        if check_input != 'y':
            sys.exit()
            
    return use_df


# =============================================================================
# ZIP FILES
# =============================================================================


def zip_files(lookup_df, base_dir):
    """
    Zip files per batch from lookup DataFrame.
    
    Args:
        lookup_df (pd.DataFrame): DataFrame with 'batch' and filename columns.
        base_dir (str): Path to folder with files (where zip files will be saved).
        include_batches (list or None): List of batch numbers to include, or None for all.
    """
    # Get logger (uses same log file as generator)
    logger = logging.getLogger()
    
    # Start timing
    start_time = datetime.now()
    logger.info(f"{'='*68}")
    logger.info("STARTING ZIP FILE CREATION")
    logger.info(f"{'='*68}")
    
    # Select only needed columns
    df = lookup_df[[BATCH_COL, FILENAME_COL]].copy()
    
    # Group filenames by batch
    grouped = df.groupby('batch')
    
    num_batches = len(grouped)
    total_files = len(df)
    logger.info(f"Found {num_batches} batch(es), {total_files} total files")
        
    batch_count = 0
    for batch, group in grouped:
        batch_count += 1
        num_rows = len(group)
        zip_filename = os.path.join(base_dir, f'Batch{batch:02d}_{num_rows}rows.zip')
        
        print(f"\nZipping batch {batch} ({num_rows} files)...")
        
        # Time this batch
        batch_start = datetime.now()
        
        # Create zip file
        with zipfile.ZipFile(zip_filename, 'w') as zipf:
            for filnavn in tqdm(group[FILENAME_COL], desc=f"Batch {batch}", unit="file"):
                file_path = os.path.join(base_dir, filnavn)
                if not os.path.exists(file_path):
                    error_msg = f"File not found: {file_path}"
                    logger.error(error_msg)
                    print(f"❌ {error_msg}")
                    raise FileNotFoundError(f"Batch {batch} aborted. File not found: {file_path}")
                zipf.write(file_path, arcname=filnavn)
        
        # Calculate batch duration
        batch_duration = (datetime.now() - batch_start).total_seconds()
        
        # Log and print batch completion with duration
        completion_msg = f"✅ Batch {batch} complete: {num_rows} files in {batch_duration/60:.1f} min to file {zip_filename}"
        logger.info(completion_msg)
        #print(completion_msg)
    
    # Calculate total duration
    end_time = datetime.now()
    duration = (end_time - start_time).total_seconds()
    
    # Final summary
    logger.info(f"{'='*68}")
    logger.info(f"✅ ZIP COMPLETE: {num_batches} batches, {total_files} files in {duration:.1f}s ({duration/60:.1f}min)")
    logger.info(f"{'='*68}")
    

# =============================================================================
# MAIN
# =============================================================================
import sys

def test_mode(args):
    print('\n\n'+'='*30 +' TEST MODE ' + '='*30)
    print(args)
    print('='*71)

    limit_rows=2

    lookup_df = load_and_validate_data(
        file_path=LOOKUP_FILE,
        required_columns=PLACEHOLDERS + [FILENAME_COL, BATCH_COL],
        limit_rows=limit_rows
    )
    
    if 'html' in args:
        print('\nLeser templatfil: ' + str(HTML_TEMPLATE_FILE))
        with open(HTML_TEMPLATE_FILE, 'r', encoding='utf-8') as f:
            template = f.read()
    else:
        print('\nLeser templatfil: ' + str(WORD_TEMPLATE_FILE))
        template = generate_html_template(WORD_TEMPLATE_FILE)

    pdf_folder = generate_pdfs(template, lookup_df)
    print('Done')
    return pdf_folder


if __name__ == "__main__":
    if len(sys.argv) > 1:
        result = test_mode(sys.argv)
        print(result)
        sys.exit()

    print('\n'+'='*100)
    print('\t\t\t\tDIGIPOST BREV-GENERATOR')
    print('='*100)
    # Get row limit from user
    limit_rows = get_row_limit()

    # Load and validate lookup data
    print('Leser lookup-fil: ' + str(LOOKUP_FILE))
    try:
        lookup_df = load_and_validate_data(
            file_path=LOOKUP_FILE,
            required_columns=PLACEHOLDERS + [FILENAME_COL, BATCH_COL],
            limit_rows=limit_rows
        )
    except ValueError as e:
        print(f"❌ Error loading or validating data: {e}")
        exit(1)
    
    print("\n" + "=" * 100)   

    # Ask: Generate new PDFs or use existing folder?
    generate_choice = input("\n❓ Generate new PDFs [G (default)] or use existing folder [E]? ").strip().lower()    
    lookup_df = select_batches(lookup_df)
    
    if generate_choice in ["g",""]:
        # Load template
        template_mode = input('\n❓ Generate template from word [W, default] or read html_template [H]? ').strip().lower()
        if template_mode in ['w','']:
            print('\nLeser templatfil: ' + str(WORD_TEMPLATE_FILE))
            template = generate_html_template(WORD_TEMPLATE_FILE)
            
        elif template_mode == 'h':
            print('\nLeser templatfil: ' + str(HTML_TEMPLATE_FILE))
            with open(HTML_TEMPLATE_FILE, 'r', encoding='utf-8') as f:
                template = f.read()

        pdf_folder = generate_pdfs(template, lookup_df)

    elif generate_choice == "e":
        root = Tk()
        root.withdraw()
        existing_folder = askdirectory(initialdir=LOOKUP_FILE.parent, title="Select folder with existing PDF files")
        root.destroy()
        
        if not existing_folder:
            print("\n❌ No folder selected")
            exit(1)

        pdf_folder = Path(existing_folder)
        print(f"\n✅ Using existing folder: {pdf_folder}")
        # SET UP LOGGING FOR EXISTING FOLDER SCENARIO
        logger = setup_logging(pdf_folder)
    
    else:
        print("\n❌ Invalid input")
        exit(1)
    
    # Ask: Create zip files?
    zip_choice = input("\n❓ Create zip files[Y/N]? ").strip().lower()
    
    if zip_choice in ["y", ""]:
        
        try:
            zip_files(
                lookup_df=lookup_df,
                base_dir=str(pdf_folder),
            )
        except Exception as e:
            print(f"\n❌ Error during zipping: {e}")
            exit(1)
    else:
        print("\n✅ Skipping zip file creation")
    
    # Done
    print("\n" + "=" * 100)
    print("✅ ALL OPERATIONS COMPLETE")
    print("=" * 100)
